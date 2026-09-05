"""
Legal Document Business Writer Tool
Comprehensive legal document management system with PDF/form handling, 
email/SMS integration for uploads and downloads.
"""

import os
import json
import smtplib
import logging
from datetime import datetime
from typing import Dict, List, Optional, Tuple
from pathlib import Path
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email.mime.text import MIMEText
from email import encoders
from abc import ABC, abstractmethod
from werkzeug.utils import secure_filename

logger = logging.getLogger(__name__)

try:
    import PyPDF2
    from PyPDF2 import PdfReader, PdfWriter
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False

try:
    import pypdfform
    FORM_SUPPORT = True
except ImportError:
    FORM_SUPPORT = False


class LegalDocumentLibrary:
    """Manages legal document templates and library."""
    
    TEMPLATES = {
        "nda": {
            "name": "Non-Disclosure Agreement",
            "category": "Agreements",
            "description": "Protects confidential business information",
            "fields": ["party_name", "disclosure_period", "restriction_period", "jurisdiction"]
        },
        "service_agreement": {
            "name": "Service Agreement",
            "category": "Agreements",
            "description": "Defines terms between service provider and client",
            "fields": ["service_description", "fee", "term", "payment_terms", "termination_clause"]
        },
        "employment_contract": {
            "name": "Employment Contract",
            "category": "Employment",
            "description": "Employment relationship agreement",
            "fields": ["employee_name", "position", "salary", "start_date", "benefits", "termination_terms"]
        },
        "independent_contractor": {
            "name": "Independent Contractor Agreement",
            "category": "Agreements",
            "description": "Contract for independent contractors",
            "fields": ["contractor_name", "scope_of_work", "rate", "project_duration", "deliverables"]
        },
        "business_proposal": {
            "name": "Business Proposal",
            "category": "Business",
            "description": "Professional business proposal template",
            "fields": ["client_name", "project_scope", "timeline", "cost_estimate", "deliverables"]
        },
        "privacy_policy": {
            "name": "Privacy Policy",
            "category": "Compliance",
            "description": "Data privacy and protection policy",
            "fields": ["company_name", "data_types", "retention_period", "third_party_sharing"]
        },
        "terms_of_service": {
            "name": "Terms of Service",
            "category": "Compliance",
            "description": "User terms and conditions",
            "fields": ["company_name", "service_description", "user_obligations", "liability_limits"]
        },
        "purchase_agreement": {
            "name": "Purchase Agreement",
            "category": "Sales",
            "description": "Agreement for sale of goods or services",
            "fields": ["seller_name", "buyer_name", "item_description", "price", "delivery_terms", "payment_method"]
        },
        "lease_agreement": {
            "name": "Lease Agreement",
            "category": "Real Estate",
            "description": "Property lease contract",
            "fields": ["landlord_name", "tenant_name", "property_address", "lease_term", "monthly_rent", "security_deposit"]
        },
        "invoice_template": {
            "name": "Professional Invoice",
            "category": "Business",
            "description": "Standard business invoice",
            "fields": ["invoice_number", "date", "client_name", "description", "amount", "due_date"]
        }
    }
    
    def __init__(self, library_path: str = "./legal_templates"):
        self.library_path = Path(library_path)
        self.library_path.mkdir(exist_ok=True)
        self.documents: Dict[str, dict] = {}
        self._initialize_templates()
    
    def _initialize_templates(self):
        """Initialize template metadata."""
        for template_id, template_data in self.TEMPLATES.items():
            self.documents[template_id] = {
                **template_data,
                "created_at": datetime.now().isoformat(),
                "file_path": self.library_path / f"{template_id}_template.docx"
            }
    
    def get_template(self, template_id: str) -> Dict:
        """Retrieve a template by ID."""
        if template_id not in self.documents:
            raise ValueError(f"Template '{template_id}' not found in library")
        return self.documents[template_id]
    
    def list_templates(self, category: Optional[str] = None) -> List[Dict]:
        """List all available templates, optionally filtered by category."""
        templates = list(self.documents.values())
        if category:
            templates = [t for t in templates if t.get("category") == category]
        return templates
    
    def list_categories(self) -> List[str]:
        """Get all available document categories."""
        return sorted(set(t.get("category") for t in self.documents.values()))


class PDFHandler:
    """Handles PDF reading, writing, and editing."""
    
    def __init__(self):
        if not PDF_SUPPORT:
            raise ImportError("PyPDF2 not installed. Install with: pip install PyPDF2")
    
    def read_pdf(self, file_path: str) -> Dict:
        """Read and extract text from PDF."""
        try:
            with open(file_path, 'rb') as file:
                reader = PdfReader(file)
                metadata = {
                    "num_pages": len(reader.pages),
                    "title": reader.metadata.title if reader.metadata else "Unknown",
                    "author": reader.metadata.author if reader.metadata else "Unknown",
                    "pages": []
                }
                
                for page_num, page in enumerate(reader.pages):
                    text = page.extract_text()
                    metadata["pages"].append({
                        "page_num": page_num + 1,
                        "content": text,
                        "num_images": len(page.images)
                    })
                
                return metadata
        except Exception as exc:
            logger.exception("Error reading PDF")
            raise IOError("Unable to read PDF") from exc
    
    def write_pdf(self, content: str, output_path: str, title: str = "Document"):
        """Create a new PDF from text content."""
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.pdfgen import canvas
            from reportlab.lib.utils import ImageReader
            
            c = canvas.Canvas(output_path, pagesize=letter)
            c.setTitle(title)
            
            # Wrap text and write to PDF
            from textwrap import wrap
            y = 750
            for line in content.split('\n'):
                wrapped_lines = wrap(line, width=80)
                for wrapped_line in wrapped_lines:
                    c.drawString(50, y, wrapped_line)
                    y -= 20
                    if y < 50:
                        c.showPage()
                        y = 750
            
            c.save()
            return {"status": "success", "path": output_path}
        except Exception as exc:
            logger.exception("Error writing PDF")
            raise IOError("Unable to write PDF") from exc
    
    def merge_pdfs(self, pdf_paths: List[str], output_path: str) -> Dict:
        """Merge multiple PDFs into one."""
        try:
            writer = PdfWriter()
            for pdf_path in pdf_paths:
                with open(pdf_path, 'rb') as f:
                    reader = PdfReader(f)
                    for page in reader.pages:
                        writer.add_page(page)
            
            with open(output_path, 'wb') as f:
                writer.write(f)
            
            return {"status": "success", "output": output_path, "num_pdfs": len(pdf_paths)}
        except Exception as exc:
            logger.exception("Error merging PDFs")
            raise IOError("Unable to merge PDFs") from exc
    
    def split_pdf(self, pdf_path: str, output_dir: str, pages: Optional[List[int]] = None) -> Dict:
        """Split PDF into individual pages or extract specific 1-based page numbers."""
        try:
            Path(output_dir).mkdir(parents=True, exist_ok=True)
            with open(pdf_path, 'rb') as f:
                reader = PdfReader(f)
                total_pages = len(reader.pages)
                
                if pages is None:
                    pages = list(range(1, total_pages + 1))
                
                output_files = []
                for page_num in pages:
                    if 1 <= page_num <= total_pages:
                        writer = PdfWriter()
                        writer.add_page(reader.pages[page_num - 1])
                        output_file = Path(output_dir) / f"page_{page_num}.pdf"
                        with open(output_file, 'wb') as out:
                            writer.write(out)
                        output_files.append(str(output_file))
                
                return {"status": "success", "output_files": output_files}
        except Exception as exc:
            logger.exception("Error splitting PDF")
            raise IOError("Unable to split PDF") from exc
    
    def add_watermark(self, pdf_path: str, watermark_text: str, output_path: str) -> Dict:
        """Add watermark to PDF."""
        try:
            from reportlab.pdfgen import canvas
            from io import BytesIO
            
            with open(pdf_path, 'rb') as f:
                reader = PdfReader(f)
                writer = PdfWriter()
                
                for page in reader.pages:
                    # Create watermark
                    watermark_buffer = BytesIO()
                    watermark_canvas = canvas.Canvas(watermark_buffer, pagesize=(page.mediabox.width, page.mediabox.height))
                    watermark_canvas.setFontSize(60)
                    watermark_canvas.setFillAlpha(0.3)
                    watermark_canvas.drawCentredString(
                        float(page.mediabox.width) / 2,
                        float(page.mediabox.height) / 2,
                        watermark_text
                    )
                    watermark_canvas.save()
                    watermark_buffer.seek(0)
                    
                    watermark_pdf = PdfReader(watermark_buffer)
                    page.merge_page(watermark_pdf.pages[0])
                    writer.add_page(page)
                
                with open(output_path, 'wb') as out:
                    writer.write(out)
            
            return {"status": "success", "output": output_path}
        except Exception as exc:
            logger.exception("Error adding PDF watermark")
            raise IOError("Unable to add watermark") from exc


class FormHandler:
    """Handles form reading and writing."""
    
    def __init__(self):
        if not FORM_SUPPORT:
            raise ImportError("pypdfform not installed. Install with: pip install pypdfform")
    
    def read_form_fields(self, pdf_path: str) -> Dict:
        """Extract form fields from a PDF form."""
        try:
            fields = pypdfform.get_form_fields(pdf_path)
            return {
                "status": "success",
                "file": pdf_path,
                "fields": fields,
                "num_fields": len(fields)
            }
        except Exception:
            logger.exception("Error reading PDF form fields")
            return {"status": "error", "message": "Unable to read form fields"}
    
    def fill_form(self, pdf_path: str, form_data: Dict, output_path: str) -> Dict:
        """Fill form fields with data."""
        try:
            filled_pdf = pypdfform.fill_form(pdf_path, form_data)
            filled_pdf.write(output_path)
            return {"status": "success", "output": output_path, "fields_filled": len(form_data)}
        except Exception:
            logger.exception("Error filling PDF form")
            return {"status": "error", "message": "Unable to fill form"}
    
    def create_form_template(self, fields: List[str], output_path: str) -> Dict:
        """Create a new form template with specified fields."""
        try:
            # This would require a PDF library that supports form creation
            # For now, return a template specification
            template = {
                "name": Path(output_path).stem,
                "fields": fields,
                "created_at": datetime.now().isoformat(),
                "description": "Form template specification"
            }
            
            with open(output_path.replace(".pdf", ".json"), 'w') as f:
                json.dump(template, f, indent=2)
            
            return {"status": "success", "template": template}
        except Exception:
            logger.exception("Error creating form template")
            return {"status": "error", "message": "Unable to create form template"}


class EmailIntegration:
    """Handles email sending and receiving with attachments."""
    
    def __init__(
        self, smtp_server: str, smtp_port: int, email: str, password: str, enable_tls: bool = True
    ):
        self.smtp_server = smtp_server
        self.smtp_port = smtp_port
        self.email = email
        self.password = password
        self.enable_tls = enable_tls
    
    def send_document(self, recipient_email: str, document_path: str, 
                     subject: str = "Legal Document", message: str = "") -> Dict:
        """Send a document via email."""
        try:
            msg = MIMEMultipart()
            msg['From'] = self.email
            msg['To'] = recipient_email
            msg['Subject'] = subject
            
            msg.attach(MIMEText(message or "Please find the attached document.", 'plain'))
            
            # Attach document
            with open(document_path, 'rb') as attachment:
                part = MIMEBase('application', 'octet-stream')
                part.set_payload(attachment.read())
                encoders.encode_base64(part)
                part.add_header('Content-Disposition', f'attachment; filename= {Path(document_path).name}')
                msg.attach(part)
            
            # SMTP port 465 uses implicit TLS; other ports can upgrade with STARTTLS.
            smtp_class = smtplib.SMTP_SSL if self.smtp_port == 465 else smtplib.SMTP
            with smtp_class(self.smtp_server, self.smtp_port) as server:
                if self.smtp_port != 465 and self.enable_tls:
                    server.ehlo()
                    if not server.has_extn("starttls"):
                        return {"status": "error", "message": "SMTP server does not support STARTTLS"}
                    server.starttls()
                    server.ehlo()
                server.login(self.email, self.password)
                server.send_message(msg)
            
            return {"status": "success", "recipient": recipient_email, "document": Path(document_path).name}
        except Exception:
            logger.exception("Error sending document email")
            return {"status": "error", "message": "Unable to send document"}
    
    def send_multiple_documents(self, recipients: List[str], document_paths: List[str],
                               subject: str = "Legal Documents", message: str = "") -> Dict:
        """Send multiple documents to multiple recipients."""
        results = []
        for recipient in recipients:
            for doc_path in document_paths:
                result = self.send_document(recipient, doc_path, subject, message)
                results.append(result)
        
        return {
            "status": "success",
            "total_sent": len(results),
            "results": results
        }
    
    def parse_document_from_email(self, email_content: str) -> Dict:
        """Extract documents from email content (attachment references)."""
        # This would integrate with email service to extract attachments
        return {
            "status": "success",
            "parsed_attachments": [],
            "message": "Email parsing requires email service integration"
        }


class SMSIntegration:
    """Handles SMS sending and receiving with document links."""
    
    def __init__(
        self,
        twilio_account_sid: str,
        twilio_auth_token: str,
        from_number: str,
        signalwire_space_url: Optional[str] = None,
    ):
        try:
            from twilio.rest import Client
            self.client = Client(twilio_account_sid, twilio_auth_token)
            if signalwire_space_url:
                space_url = signalwire_space_url.strip().rstrip("/")
                if not space_url.startswith("http"):
                    space_url = f"https://{space_url}"
                self.client.api.base_url = space_url
            self.from_number = from_number
            self.sms_support = True
        except ImportError:
            self.sms_support = False
    
    def send_document_link(self, recipient_number: str, document_url: str,
                          message: str = "Your document is ready") -> Dict:
        """Send document link via SMS."""
        if not self.sms_support:
            return {"status": "error", "message": "Twilio not installed"}
        
        try:
            full_message = f"{message}\n\nDownload: {document_url}"
            message_obj = self.client.messages.create(
                body=full_message,
                from_=self.from_number,
                to=recipient_number
            )
            
            return {
                "status": "success",
                "message_id": message_obj.sid,
                "recipient": recipient_number
            }
        except Exception:
            logger.exception("Error sending document SMS")
            return {"status": "error", "message": "Unable to send document"}
    
class LegalDocumentWriter:
    """Main class for legal document generation and management."""

    SUPPORTED_OUTPUT_FORMATS = {"docx", "pdf", "txt"}

    def __init__(self, config: Optional[Dict] = None):
        self.config = config or {}
        self.library = LegalDocumentLibrary(self.config.get("library_path", "./legal_templates"))
        self.pdf_handler = PDFHandler() if PDF_SUPPORT else None
        self.form_handler = FormHandler() if FORM_SUPPORT else None
        self.email_integration = None
        self.sms_integration = None
        self.documents_dir = Path(self.config.get("documents_dir", "./generated_documents")).resolve()
        self.documents_dir.mkdir(parents=True, exist_ok=True)

    def resolve_document_path(self, file_path: str, require_exists: bool = True) -> Path:
        """Return a file path only when it remains inside managed document storage."""
        safe_filename = secure_filename(os.path.basename(str(file_path)))
        if not safe_filename:
            raise ValueError("A valid document filename is required")

        resolved_path = (self.documents_dir / safe_filename).resolve()
        try:
            resolved_path.relative_to(self.documents_dir)
        except ValueError as exc:
            raise ValueError("Document path must be inside managed document storage") from exc

        if require_exists and not resolved_path.is_file():
            raise FileNotFoundError("Document not found")

        return resolved_path

    def create_output_path(self, filename: str, format_type: str) -> Path:
        """Create a sanitized path for a generated document."""
        normalized_format = str(format_type).lower()
        if normalized_format not in self.SUPPORTED_OUTPUT_FORMATS:
            raise ValueError("Unsupported document format")

        safe_filename = secure_filename(
            os.path.splitext(os.path.basename(str(filename)))[0]
        )
        if not safe_filename:
            raise ValueError("A valid document filename is required")

        return self.resolve_document_path(
            f"{safe_filename}.{normalized_format}",
            require_exists=False,
        )

    def setup_email(
        self, smtp_server: str, smtp_port: int, email: str, password: str, enable_tls: bool = True
    ):
        """Setup email integration."""
        self.email_integration = EmailIntegration(
            smtp_server, smtp_port, email, password, enable_tls
        )
    
    def setup_sms(
        self,
        twilio_account_sid: str,
        twilio_auth_token: str,
        from_number: str,
        signalwire_space_url: Optional[str] = None,
    ):
        """Set up Twilio or SignalWire document-link delivery."""
        self.sms_integration = SMSIntegration(
            twilio_account_sid,
            twilio_auth_token,
            from_number,
            signalwire_space_url,
        )
    
    def generate_document(self, template_id: str, data: Dict,
                         format: str = "docx", filename: Optional[str] = None) -> Dict:
        """Generate a document from a template with provided data."""
        try:
            template = self.library.get_template(template_id)

            format = str(format).lower()
            if format not in self.SUPPORTED_OUTPUT_FORMATS:
                return {"status": "error", "message": "Unsupported document format"}
            if format == "docx" and not DOCX_SUPPORT:
                return {
                    "status": "error",
                    "message": "DOCX output is unavailable because python-docx is not installed",
                }
            if format == "pdf" and not self.pdf_handler:
                return {
                    "status": "error",
                    "message": "PDF output is unavailable because PDF support is not installed",
                }

            if filename is None:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = f"{template_id}_{timestamp}"
            
            output_path = self.create_output_path(filename, format)
            
            # Generate document content
            content = self._render_template(template, data)
            
            if format == "docx" and DOCX_SUPPORT:
                self._write_docx(content, output_path)
            elif format == "pdf" and self.pdf_handler:
                self.pdf_handler.write_pdf(content, str(output_path), title=template["name"])
            elif format == "txt":
                with open(output_path, 'w') as f:
                    f.write(content)
            else:
                raise ValueError("Unsupported document format")
            
            return {
                "status": "success",
                "template": template_id,
                "file": str(output_path),
                "format": format,
                "size": output_path.stat().st_size
            }
        except Exception:
            logger.exception("Error generating document")
            return {"status": "error", "message": "Unable to generate document"}
    
    def _render_template(self, template: Dict, data: Dict) -> str:
        """Render template with data."""
        content = f"# {template['name']}\n\n"
        content += f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n"
        
        for field in template.get("fields", []):
            field_label = field.replace("_", " ").title()
            field_value = data.get(field, "[TO BE FILLED]")
            content += f"{field_label}: {field_value}\n\n"
        
        return content
    
    def _write_docx(self, content: str, output_path: Path):
        """Write content to DOCX file."""
        doc = Document()
        for line in content.split('\n'):
            if line.startswith('#'):
                p = doc.add_paragraph(line.lstrip('#').strip(), style='Heading 1')
            elif line.strip():
                doc.add_paragraph(line)
        doc.save(str(output_path))
    
    def export_document(self, file_path: str, format: str = "pdf") -> Dict:
        """Export document to different format."""
        try:
            source_path = self.resolve_document_path(file_path)
            output_path = self.create_output_path(source_path.stem, format)
            # Conversion logic would go here
            
            return {
                "status": "success",
                "original": str(source_path),
                "converted": str(output_path),
                "format": format
            }
        except Exception:
            logger.exception("Error exporting document")
            return {"status": "error", "message": "Unable to export document"}
    
    def email_document(self, file_path: str, recipient: str, subject: str = "Document",
                      message: str = "") -> Dict:
        """Email a document to recipient."""
        if not self.email_integration:
            return {"status": "error", "message": "Email not configured"}
        
        document_path = self.resolve_document_path(file_path)
        return self.email_integration.send_document(
            recipient, str(document_path), subject, message
        )
    
    def sms_document(
        self, file_path: str, recipient_number: str, document_url: str
    ) -> Dict:
        """Send an SMS containing a secure document-download link."""
        if not self.sms_integration:
            return {"status": "error", "message": "SMS not configured"}

        self.resolve_document_path(file_path)
        return self.sms_integration.send_document_link(recipient_number, document_url)
    
    def list_generated_documents(self) -> List[Dict]:
        """List all generated documents."""
        docs = []
        for doc_path in self.documents_dir.iterdir():
            if doc_path.is_file():
                docs.append({
                    "name": doc_path.name,
                    "path": str(doc_path),
                    "size": doc_path.stat().st_size,
                    "created": datetime.fromtimestamp(doc_path.stat().st_ctime).isoformat()
                })
        return sorted(docs, key=lambda x: x["created"], reverse=True)


if __name__ == "__main__":
    # Example usage
    writer = LegalDocumentWriter()
    
    # List available templates
    print("\n=== Available Legal Document Templates ===")
    templates = writer.library.list_templates()
    for template in templates:
        print(f"  • {template['name']} ({template['category']})")
    
    # Generate a sample document
    print("\n=== Generating Sample NDA ===")
    result = writer.generate_document(
        "nda",
        {
            "party_name": "Acme Corporation",
            "disclosure_period": "3 years",
            "restriction_period": "5 years",
            "jurisdiction": "California"
        },
        format="docx"
    )
    print(f"  Generated: {result['file']}")
